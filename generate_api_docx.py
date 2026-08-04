# -*- coding: utf-8 -*-
"""从 OpenAPI/Swagger 生成 CRM-OKR System API 中文 Word 文档。"""
from __future__ import annotations

import json
import os
import sys
import urllib.request
from collections import defaultdict
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_PROJECT = Path(r"C:\Users\Administrator\Desktop\crm-okr-system\CRM-OKR-API文档.docx")
OUT_DESKTOP = Path(r"C:\Users\Administrator\Desktop\CRM-OKR-API文档.docx")
OPENAPI_URL = "http://127.0.0.1:8000/openapi.json"
OPENAPI_CACHE = Path(r"C:\Users\Administrator\Desktop\crm-okr-system\openapi.json")

METHOD_COLORS = {
    "GET": RGBColor(0x22, 0x8B, 0x22),
    "POST": RGBColor(0x1E, 0x90, 0xFF),
    "PUT": RGBColor(0xE6, 0xA8, 0x17),
    "PATCH": RGBColor(0x99, 0x32, 0xCC),
    "DELETE": RGBColor(0xDC, 0x14, 0x3C),
    "HEAD": RGBColor(0x70, 0x80, 0x90),
    "OPTIONS": RGBColor(0x70, 0x80, 0x90),
}


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)


def add_para(doc, text, *, size=11, bold=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading_cn(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 16, 2: 14, 3: 12}.get(level, 11), bold=True)
    return p


def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, size=9, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text) if text is not None else "")
    set_run_font(run, size=size, bold=bold)
    if fill:
        set_cell_shading(cell, fill)
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, fill="D9E2F3")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=9)
    doc.add_paragraph()
    return table


def load_openapi() -> dict:
    if OPENAPI_CACHE.exists() and OPENAPI_CACHE.stat().st_size > 1000:
        print(f"Using cached {OPENAPI_CACHE}")
        return json.loads(OPENAPI_CACHE.read_text(encoding="utf-8"))

    try:
        print(f"Downloading {OPENAPI_URL} ...")
        with urllib.request.urlopen(OPENAPI_URL, timeout=30) as resp:
            data = resp.read()
        OPENAPI_CACHE.write_bytes(data)
        print(f"Saved {OPENAPI_CACHE} ({len(data)} bytes)")
        return json.loads(data.decode("utf-8"))
    except Exception as e:
        print(f"Download failed: {e}")
        print("Falling back to app.openapi() ...")
        backend = Path(__file__).resolve().parent / "backend"
        sys.path.insert(0, str(backend))
        os.chdir(backend)
        from app.main import app  # noqa: WPS433

        spec = app.openapi()
        OPENAPI_CACHE.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")
        return spec


def resolve_ref(spec: dict, ref: str) -> dict:
    if not ref or not ref.startswith("#/"):
        return {}
    node: dict | list | str | None = spec
    for part in ref[2:].split("/"):
        if not isinstance(node, dict):
            return {}
        node = node.get(part)
    return deepcopy(node) if isinstance(node, dict) else {}


def schema_type_label(schema: dict | None, spec: dict) -> str:
    if not schema:
        return "-"
    if "$ref" in schema:
        return schema["$ref"].split("/")[-1]
    if "allOf" in schema and schema["allOf"]:
        parts = [schema_type_label(s, spec) for s in schema["allOf"]]
        return " & ".join(p for p in parts if p != "-") or "object"
    if "anyOf" in schema:
        non_null = [s for s in schema["anyOf"] if s.get("type") != "null" and s.get("type") != None]
        if len(non_null) == 1:
            base = schema_type_label(non_null[0], spec)
            return f"{base}?" if len(schema["anyOf"]) > 1 else base
        return " | ".join(schema_type_label(s, spec) for s in schema["anyOf"])
    if "oneOf" in schema:
        return " | ".join(schema_type_label(s, spec) for s in schema["oneOf"])
    t = schema.get("type", "object")
    if t == "array":
        return f"array[{schema_type_label(schema.get('items', {}), spec)}]"
    if schema.get("format"):
        return f"{t}({schema['format']})"
    if schema.get("enum"):
        return f"enum({', '.join(map(str, schema['enum'][:6]))}{'...' if len(schema['enum']) > 6 else ''})"
    return str(t)


def summarize_schema_fields(schema: dict, spec: dict, limit: int = 20) -> list[tuple[str, str, str]]:
    """Return list of (name, type, required_flag)."""
    schema = deepcopy(schema)
    while "$ref" in schema:
        schema = resolve_ref(spec, schema["$ref"])
    if "allOf" in schema:
        merged: dict = {"properties": {}, "required": []}
        for part in schema["allOf"]:
            part = resolve_ref(spec, part["$ref"]) if "$ref" in part else part
            merged["properties"].update(part.get("properties") or {})
            merged["required"].extend(part.get("required") or [])
        schema = merged
    props = schema.get("properties") or {}
    required = set(schema.get("required") or [])
    rows = []
    for name, prop in list(props.items())[:limit]:
        rows.append((name, schema_type_label(prop, spec), "是" if name in required else "否"))
    if len(props) > limit:
        rows.append(("...", f"另有 {len(props) - limit} 个字段", ""))
    return rows


def request_body_summary(op: dict, spec: dict) -> str:
    body = op.get("requestBody")
    if not body:
        return "无"
    content = body.get("content") or {}
    parts = []
    for ctype, media in content.items():
        schema = media.get("schema") or {}
        label = schema_type_label(schema, spec)
        req = "必填" if body.get("required") else "可选"
        parts.append(f"{ctype} → {label}（{req}）")
        fields = summarize_schema_fields(
            resolve_ref(spec, schema["$ref"]) if "$ref" in schema else schema,
            spec,
            limit=12,
        )
        if fields:
            field_txt = ", ".join(f"{n}:{t}" for n, t, _ in fields[:8])
            parts.append(f"  主要字段: {field_txt}")
    return "\n".join(parts) if parts else "无"


def collect_endpoints(spec: dict):
    grouped: dict[str, list] = defaultdict(list)
    paths = spec.get("paths") or {}
    for path, methods in paths.items():
        for method, op in methods.items():
            if method.startswith("x-") or not isinstance(op, dict):
                continue
            if method.lower() not in ("get", "post", "put", "delete", "patch", "head", "options"):
                continue
            tags = op.get("tags") or ["未分类"]
            tag = tags[0]
            grouped[tag].append(
                {
                    "method": method.upper(),
                    "path": path,
                    "summary": op.get("summary") or "",
                    "description": op.get("description") or "",
                    "parameters": op.get("parameters") or [],
                    "operation": op,
                }
            )
    # stable order: by tag name, then path, then method
    for tag in grouped:
        grouped[tag].sort(key=lambda x: (x["path"], x["method"]))
    return dict(sorted(grouped.items(), key=lambda kv: kv[0]))


def param_type(param: dict, spec: dict) -> str:
    if "schema" in param:
        return schema_type_label(param["schema"], spec)
    if "content" in param:
        for media in (param["content"] or {}).values():
            return schema_type_label(media.get("schema"), spec)
    return "-"


def build_document(spec: dict) -> Document:
    info = spec.get("info") or {}
    title = info.get("title") or "CRM-OKR System API"
    version = info.get("version") or "-"
    description = info.get("description") or ""

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ----- Title page -----
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "CRM-OKR System API 文档", size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, title, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, f"版本：{version}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    if description:
        add_para(doc, description, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    paths_count = len(spec.get("paths") or {})
    schemas_count = len((spec.get("components") or {}).get("schemas") or {})
    endpoint_count = sum(
        1
        for p in (spec.get("paths") or {}).values()
        for m, op in p.items()
        if m.lower() in ("get", "post", "put", "delete", "patch", "head", "options") and isinstance(op, dict)
    )
    add_para(
        doc,
        f"接口路径数：{paths_count}　|　端点数：{endpoint_count}　|　数据模型：{schemas_count}",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_para(doc, "由 OpenAPI / Swagger 自动生成", size=10, color=RGBColor(0x66, 0x66, 0x66), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    grouped = collect_endpoints(spec)

    # ----- TOC style -----
    add_heading_cn(doc, "目录（按模块分组）", level=1)
    add_para(doc, "以下按 OpenAPI tags 分组列出各模块及接口数量：", size=10, color=RGBColor(0x55, 0x55, 0x55))
    toc_rows = [(tag, str(len(items))) for tag, items in grouped.items()]
    add_table(doc, ["模块（Tag）", "接口数"], toc_rows)
    doc.add_page_break()

    # ----- Endpoints -----
    add_heading_cn(doc, "接口详情", level=1)
    for tag, items in grouped.items():
        add_heading_cn(doc, f"{tag}（{len(items)}）", level=2)
        for ep in items:
            method = ep["method"]
            color = METHOD_COLORS.get(method, RGBColor(0, 0, 0))
            # method + path heading
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            r1 = p.add_run(method)
            set_run_font(r1, size=12, bold=True, color=color)
            r2 = p.add_run(f"  {ep['path']}")
            set_run_font(r2, size=12, bold=True)

            if ep["summary"]:
                add_para(doc, f"摘要：{ep['summary']}", size=10, space_after=2)
            if ep["description"] and ep["description"] != ep["summary"]:
                desc = ep["description"].strip()
                if len(desc) > 800:
                    desc = desc[:800] + "…"
                add_para(doc, f"说明：{desc}", size=10, space_after=4)

            # parameters
            params = ep["parameters"]
            # also resolve path-level params already merged in OpenAPI usually
            if params:
                add_para(doc, "请求参数：", size=10, bold=True, space_after=2)
                rows = []
                for prm in params:
                    # resolve $ref parameters
                    if "$ref" in prm:
                        prm = resolve_ref(spec, prm["$ref"])
                    rows.append(
                        (
                            prm.get("name", ""),
                            prm.get("in", ""),
                            "是" if prm.get("required") else "否",
                            param_type(prm, spec),
                            (prm.get("description") or "")[:120],
                        )
                    )
                add_table(doc, ["名称", "位置", "必填", "类型", "说明"], rows)
            else:
                add_para(doc, "请求参数：无", size=10, space_after=2)

            # request body
            body_txt = request_body_summary(ep["operation"], spec)
            add_para(doc, "请求体：", size=10, bold=True, space_after=2)
            for line in body_txt.split("\n"):
                add_para(doc, line, size=9, space_after=1)

            # responses
            responses = ep["operation"].get("responses") or {}
            if responses:
                add_para(doc, "响应码：", size=10, bold=True, space_after=2)
                resp_rows = []
                for code, resp in responses.items():
                    if isinstance(resp, dict) and "$ref" in resp:
                        resp = resolve_ref(spec, resp["$ref"])
                    desc = (resp or {}).get("description") or ""
                    schema_lbl = ""
                    content = (resp or {}).get("content") or {}
                    for media in content.values():
                        schema_lbl = schema_type_label(media.get("schema"), spec)
                        break
                    resp_rows.append((str(code), desc[:100], schema_lbl or "-"))
                add_table(doc, ["状态码", "说明", "响应模型"], resp_rows)

    # ----- Schemas -----
    doc.add_page_break()
    add_heading_cn(doc, "数据模型（Schemas）", level=1)
    schemas = (spec.get("components") or {}).get("schemas") or {}
    add_para(doc, f"共 {len(schemas)} 个数据模型，下列列出各模型关键字段：", size=10, space_after=8)

    for name in sorted(schemas.keys()):
        schema = schemas[name]
        add_heading_cn(doc, name, level=3)
        stype = schema.get("type", "object")
        if schema.get("description"):
            add_para(doc, schema["description"][:300], size=9, color=RGBColor(0x55, 0x55, 0x55), space_after=2)
        if schema.get("enum"):
            add_para(doc, f"类型：枚举　取值：{', '.join(map(str, schema['enum']))}", size=9, space_after=4)
            continue
        fields = summarize_schema_fields(schema, spec, limit=25)
        if fields:
            add_table(doc, ["字段", "类型", "必填"], fields)
        else:
            add_para(doc, f"类型：{stype}（无展开字段）", size=9, space_after=4)

    return doc


def main():
    try:
        import docx  # noqa: F401
    except ImportError:
        import subprocess

        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])

    spec = load_openapi()
    doc = build_document(spec)
    OUT_PROJECT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PROJECT))
    doc.save(str(OUT_DESKTOP))

    size = OUT_PROJECT.stat().st_size
    paths_count = len(spec.get("paths") or {})
    schemas_count = len((spec.get("components") or {}).get("schemas") or {})
    endpoint_count = sum(
        1
        for p in (spec.get("paths") or {}).values()
        for m, op in p.items()
        if m.lower() in ("get", "post", "put", "delete", "patch", "head", "options") and isinstance(op, dict)
    )
    print("OK")
    print(f"FILE={OUT_PROJECT}")
    print(f"DESKTOP={OUT_DESKTOP}")
    print(f"PATHS={paths_count}")
    print(f"ENDPOINTS={endpoint_count}")
    print(f"SCHEMAS={schemas_count}")
    print(f"SIZE={size}")
    print(f"SIZE_KB={size / 1024:.1f}")


if __name__ == "__main__":
    main()
