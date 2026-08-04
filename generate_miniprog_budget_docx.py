# -*- coding: utf-8 -*-
"""生成微信小程序预算清单 Word 文档"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from pathlib import Path


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if fill:
        set_cell_shading(cell, fill)
    # 垂直居中
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


def add_heading_cn(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
        p.space_before = Pt(14)
        p.space_after = Pt(8)
    else:
        run = p.add_run(text)
        set_run_font(run, size=13, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))
        p.space_before = Pt(10)
        p.space_after = Pt(6)
    return p


def add_para(doc, text, size=11, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def fill_table(table, rows, header=True):
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            is_header = header and i == 0
            set_cell_text(
                table.rows[i].cells[j],
                str(val),
                bold=is_header,
                size=10,
                align=WD_ALIGN_PARAGRAPH.CENTER if j != 1 and j != 4 else WD_ALIGN_PARAGRAPH.LEFT,
                fill="1F4E79" if is_header else ("F2F2F2" if i % 2 == 0 else None),
            )
            if is_header:
                for p in table.rows[i].cells[j].paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def set_table_width(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)


def main():
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("微信小程序项目预算清单")
    set_run_font(run, size=22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("基础设施与平台资质费用估算（首年）")
    set_run_font(run, size=12, color=RGBColor(0x66, 0x66, 0x66))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("编制日期：2026年8月3日  |  币种：人民币（元）")
    set_run_font(run, size=10, color=RGBColor(0x88, 0x88, 0x88))

    add_heading_cn(doc, "一、说明", 1)
    add_para(
        doc,
        "本清单用于估算微信小程序正式上线所需的平台资质、域名与云基础设施等费用，不含开发人力、设计外包及业务营销费用。"
        "金额为市场常见参考价，实际以官方报价及云厂商活动价为准。",
    )
    add_para(
        doc,
        "特别说明：微信小程序认证与企业微信认证为两项独立费用，是否同时发生取决于业务是否对接企业微信。",
        size=10,
    )

    add_heading_cn(doc, "二、平台与资质费用", 1)
    rows1 = [
        ["序号", "费用项", "参考金额", "周期", "说明"],
        ["1", "微信小程序企业认证", "300", "年", "企业主体正式上线通常必需；审核无论成败费用不退"],
        ["2", "企业微信认证（可选）", "300", "年", "成员规模≤1000人；中型约3000、大型约30000"],
        ["3", "域名（.com/.cn）", "50～100", "年", "小程序业务域名、服务器域名需配置"],
        ["4", "ICP备案", "0", "一次性", "国内服务器必须备案，周期约1～3周"],
        ["5", "SSL证书", "0", "年", "可使用免费DV证书"],
        ["", "小计（仅小程序，不含企微）", "约350～400", "年", "认证300 + 域名约80"],
        ["", "小计（含企业微信认证）", "约650～700", "年", "在上项基础上增加企微认证300"],
    ]
    t1 = doc.add_table(rows=len(rows1), cols=5)
    t1.style = "Table Grid"
    fill_table(t1, rows1)
    set_table_width(t1)

    add_heading_cn(doc, "三、服务器 / 后端费用（三档方案）", 1)
    add_para(doc, "建议按业务规模择一，预算表中可写“推荐档 + 活动价备注”。", size=10)

    rows2 = [
        ["方案", "配置思路", "首年参考", "适用场景", "备注"],
        ["A. 微信云开发", "基础套餐约19.9元/月", "约240", "功能简单、希望少运维", "超出配额按量计费"],
        ["B. 轻量云服务器", "2核4G / 4核4G", "约100～300", "自建API+数据库（推荐起步）", "活动价更低，续费可能上涨"],
        ["C. 正式业务机", "2核4G～4核8G+备份", "约500～1500", "用户较多、要求稳定备份", "偏企业/内部系统更合适"],
    ]
    t2 = doc.add_table(rows=len(rows2), cols=5)
    t2.style = "Table Grid"
    fill_table(t2, rows2)
    set_table_width(t2)

    add_para(
        doc,
        "实务建议：编制预算时服务器建议按 300～800 元/年 预留（含一定余量）；秒杀活动价可写备注，不宜单独作为批复依据。",
        size=10,
    )

    add_heading_cn(doc, "四、可选增值费用", 1)
    rows3 = [
        ["序号", "费用项", "参考金额", "周期", "说明"],
        ["1", "对象存储 / CDN", "0～200", "年", "图片、静态资源较多时启用"],
        ["2", "短信验证码", "约0.03～0.05元/条", "按量", "按实际发送量结算"],
        ["3", "微信支付商户", "开通免费", "—", "有交易时手续费约0.6%"],
        ["4", "预留机动费用", "200", "年", "流量突增、临时扩容等"],
    ]
    t3 = doc.add_table(rows=len(rows3), cols=5)
    t3.style = "Table Grid"
    fill_table(t3, rows3)
    set_table_width(t3)

    add_heading_cn(doc, "五、首年合计建议（可直接提交）", 1)

    rows4 = [
        ["场景", "费用构成", "建议预算", "说明"],
        [
            "场景一：纯微信小程序",
            "小程序认证 + 域名 + 轻量服务器",
            "700～1,200",
            "不含开发人力；推荐方案B",
        ],
        [
            "场景二：小程序 + 企业微信",
            "场景一 + 企微认证300",
            "1,000～1,500",
            "需企微工作台或企微对接时选用",
        ],
    ]
    t4 = doc.add_table(rows=len(rows4), cols=4)
    t4.style = "Table Grid"
    # 自定义填充：最后一列左对齐
    for i, row_data in enumerate(rows4):
        for j, val in enumerate(row_data):
            is_header = i == 0
            set_cell_text(
                t4.rows[i].cells[j],
                str(val),
                bold=is_header or j == 2,
                size=10,
                align=WD_ALIGN_PARAGRAPH.CENTER if j in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT,
                fill="1F4E79" if is_header else ("E2EFDA" if i == 1 else "FFF2CC"),
            )
            if is_header:
                for p in t4.rows[i].cells[j].paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_table_width(t4)

    add_heading_cn(doc, "六、明细示例（场景一 · 推荐写法）", 1)

    rows5 = [
        ["类别", "项目", "金额（元）", "备注"],
        ["平台与资质", "微信小程序企业认证", "300", "年费"],
        ["平台与资质", "域名", "80", "按实际注册价"],
        ["基础设施", "云服务器（2核4G轻量）", "300", "按稳妥价，非秒杀价"],
        ["基础设施", "SSL / ICP备案", "0", "免费证书 + 免费备案"],
        ["可选预留", "短信 / CDN / 机动", "200", "按需启用"],
        ["合计", "首年基础设施建议预算", "880", "约700～1,200区间内"],
    ]
    t5 = doc.add_table(rows=len(rows5), cols=4)
    t5.style = "Table Grid"
    for i, row_data in enumerate(rows5):
        for j, val in enumerate(row_data):
            is_header = i == 0
            is_total = i == len(rows5) - 1
            set_cell_text(
                t5.rows[i].cells[j],
                str(val),
                bold=is_header or is_total,
                size=10,
                align=WD_ALIGN_PARAGRAPH.CENTER if j in (2,) else WD_ALIGN_PARAGRAPH.LEFT,
                fill="1F4E79" if is_header else ("D6EAF8" if is_total else None),
            )
            if is_header:
                for p in t5.rows[i].cells[j].paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_table_width(t5)

    add_heading_cn(doc, "七、注意事项", 1)
    notes = [
        "认证费用为第三方审核服务费，认证失败一般不予退还，请一次性备齐营业执照、对公账户/法人信息等材料。",
        "微信小程序认证、公众号认证、企业微信认证相互独立，按实际开通项分别计费。",
        "云服务器活动价（如几十元/年）仅作采购参考；批复预算建议按常规续费水平预留，避免次年费用翻倍被动。",
        "本清单不含：产品设计、开发实施、测试验收、运维外包、推广获客等人力或服务费用；如需可另附「实施费用表」。",
        "涉及支付、用户隐私、行业资质（医疗/金融等）时，可能另有合规与商户审核成本，需单独评估。",
    ]
    for idx, note in enumerate(notes, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {note}")
        set_run_font(run, size=10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.35

    add_heading_cn(doc, "八、审批栏", 1)
    rows6 = [
        ["角色", "姓名", "意见", "签字", "日期"],
        ["编制人", "", "", "", ""],
        ["部门负责人", "", "", "", ""],
        ["财务审核", "", "", "", ""],
        ["批准人", "", "", "", ""],
    ]
    t6 = doc.add_table(rows=len(rows6), cols=5)
    t6.style = "Table Grid"
    for i, row_data in enumerate(rows6):
        for j, val in enumerate(row_data):
            is_header = i == 0
            set_cell_text(
                t6.rows[i].cells[j],
                str(val) if val else "　",
                bold=is_header,
                size=10,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                fill="1F4E79" if is_header else None,
            )
            if is_header:
                for p in t6.rows[i].cells[j].paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # 增高签名行
            if not is_header:
                t6.rows[i].height = Cm(1.2)
    set_table_width(t6)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("\n（本文件为费用估算稿，最终以实际采购合同及官方账单为准）")
    set_run_font(run, size=9, color=RGBColor(0x99, 0x99, 0x99))

    out = Path(r"C:\Users\Administrator\Desktop\微信小程序预算清单.docx")
    doc.save(str(out))
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
